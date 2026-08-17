"""
Word (.docx) 文本提取工具
===================
pip install python-docx
python scripts/extract_docx.py <docx路径> [输出txt路径]
"""

import sys
import os

try:
    import docx
except ImportError:
    print("请先安装 python-docx: pip install python-docx")
    sys.exit(1)


def extract_docx(docx_path: str) -> str:
    document = docx.Document(docx_path)

    parts = []
    for p in document.paragraphs:
        text = p.text.strip()
        if text:
            parts.append(text)

    # 表格内容也提取，按「单元格 | 单元格」拼成一行
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n".join(parts)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python scripts/extract_docx.py <docx路径> [输出txt路径]")
        sys.exit(1)

    docx_path = sys.argv[1]
    if not os.path.exists(docx_path):
        print(f"文件不存在: {docx_path}")
        sys.exit(1)

    output_path = sys.argv[2] if len(sys.argv) > 2 else docx_path.replace(".docx", ".txt")

    text = extract_docx(docx_path)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(text)

    print(f"提取完成: {len(text)} 字符 -> {output_path}")
